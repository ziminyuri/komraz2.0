from flask import Flask
from flaskext.mysql import MySQL  # pip install flask-mysql
from flask import Flask, render_template, request, url_for, redirect
import hashlib
from flask import send_file
from docx import Document
import re

mysql = MySQL()
app = Flask(__name__)

# Session of users
user_agent = "0"
user_sales = "0"
# user_analitics = "0"
user_admin = "0"

# MySQL configurations
# Необходимо править в соответсвии с настройками своей локальной базы
app.config['MYSQL_DATABASE_USER'] = 'bd1'
app.config['MYSQL_DATABASE_PASSWORD'] = '1'
app.config['MYSQL_DATABASE_DB'] = 'komraz8'
app.config['MYSQL_DATABASE_HOST'] = '192.168.64.2'
mysql.init_app(app)


# Функция получения хеша
def get_hash(password):
    hash = hashlib.sha256()
    hash.update(password.encode('utf-8'))
    password_hash = hash.hexdigest()
    return password_hash

def sort_rent(rent):
    sort_rental = []
    temp_primary_id = []
    for row in rent:
        temp_primary_id.append(row[0])
    temp_primary_id = sorted(temp_primary_id)
    for temp in temp_primary_id:
        for row in rent:
            if (row[0]==temp):
                sort_rental.append(row)
    return sort_rental

@app.route('/', methods=['GET', 'POST'])
def main():
    el = " "  # Ошибки на форму
    global user_admin
    global user_agent
    global user_sales
    user_admin = "0"
    user_agent = "0"
    user_sales = "0"

    if request.method == 'POST':
        try:
            username = request.form['username']
            pas = request.form['pas']
            type = request.form['type']
            conn = mysql.connect()
            cursor = conn.cursor()

            cursor.callproc('login')
            re = cursor.fetchall()

            cursor.execute("SELECT * FROM Staff")
            staff = cursor.fetchall()

            cursor.close()
            conn.close()
            for row in re:
                if username == row[0]:
                    pas = get_hash(pas)

                    if pas == row[1]:

                        if row[2] == "Агент":
                            if type == "1":
                                user_agent = str(row[3])
                                return redirect(url_for('agentPage'))
                        else:
                            if row[2] == "Сотрудник отдела недвижимости":
                                if type == "2":
                                    user_sales = str(row[3])
                                    return redirect(url_for('salesPage'))
                            else:

                                el = "Не существует такого пользователя"
                                pass
                    else:
                        el = "Неверная пара пароль/логин"
                        pass
                else:
                    el = "Неверная пара пароль/логин"
                    pass

        except Exception as e:
            print("Поподаем в exception")
            print(e)

    return render_template('index.html', e=el)


@app.route("/agentPage", methods=['GET', 'POST'])
def agentPage():
    return render_template('Agent/agent.html')


@app.route("/salesPage", methods=['GET', 'POST'])
def salesPage():
    return render_template('Sales/salesPage.html')


######################################################################################
#########################          Справочники         ###############################
######################################################################################

# Cтраница просмотров физических лиц
@app.route("/fio", methods=['GET', 'POST'])
def show_fio():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM Natural_person")
        re = cursor.fetchall()
        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return render_template('fio.html', fio=re,agent_flag = user_agent)


# Cтраница просмотров юридических лиц
@app.route("/legal_entities", methods=['GET', 'POST'])
def show_legal_entities():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM Legal_entity")

        re = cursor.fetchall()

        print(re)
        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return render_template('legal_entities.html', legal_entities=re,agent_flag = user_agent)


# Cтраница просмотров периодов оплаты
@app.route("/payment_period", methods=['GET', 'POST'])
def show_payment_period():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM Payment_period")

        re = cursor.fetchall()
        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return render_template('payment_period.html', payment_period=re,agent_flag = user_agent)


# Cтраница просмотров виды договоров
@app.route("/types_contracts", methods=['GET', 'POST'])
def show_types_contracts():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_type_contract')

        re = cursor.fetchall()
        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return render_template('types_contracts.html', types_contracts=re,agent_flag = user_agent)


# Cтраница просмотров виды услуг
@app.route("/type_service", methods=['GET', 'POST'])
def show_type_service():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM Type_of_service")

        re = cursor.fetchall()
        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return render_template('type_service.html', type_service=re,agent_flag = user_agent)


######################################################################################
#########################          Договоры         ##################################
######################################################################################

# Cтраница просмотра договоров аренды
@app.route("/lease_contract", methods=['GET', 'POST'])
def show_lease_contract():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_lease_contract')

        r = cursor.fetchall()

        cursor.callproc('show_type_contract_of_rent')
        t = cursor.fetchall()


        cursor.execute("SELECT * FROM Payment_period")
        p = cursor.fetchall()


    ###Новая конструкция
        cursor.execute("SELECT * FROM Contract")
        contract = cursor.fetchall()

        cursor.execute("SELECT * FROM Search_object")
        search_object = cursor.fetchall()

        cursor.execute("SELECT * FROM Rental_object")
        rental = cursor.fetchall()

        contact_rent = []
        contact_search = []



        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return render_template('lease_contract.html', lease_contract=r, object=rental, type_contract=t, period=p,
                           search_object=search_object)


######################################################################################
#########################          Договоры: Агенты          #########################
######################################################################################


# Cтраница просмотра договоров оказания услуг
@app.route("/rent_contract_agent", methods=['GET', 'POST'])
def show_rent_contract1():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_lease_contract')
        lease_contracts = cursor.fetchall()

        cursor.callproc('show_rental_object')
        rental_objects = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        clients = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        legal_entitys = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        natural_persons = cursor.fetchall()

        cursor.callproc('show_search_object')
        search_objects = cursor.fetchall()

        all_lease_contract = []

        '''
        Состав процедуры  - show_lease_contract
        
        Contract.Number_of_contract    - 0
        Contract.Date                  - 1
        Contract.Amount                - 2
        Rental_object.Address          - 3
        Type_contract.Title            - 4
        Kind_of_contract.Title         - 5
        Contract.ID_staff              - 6
        Contract.ID_rental_object      - 7
        Contract.ID_search_object      - 8
        Rental_object.Square           - 9
        Rental_object.Description      - 10
        Rental_object.Year_built       - 11
        Rental_object.Wall_material    - 12
        Rental_object.Floor            - 13
        Rental_object.Number_of_floors - 14
        
        '''
        '''
        
        Состав процедуры  - show_rental_object
        + отмечены поля, которые имеются в show_lease_contract
        
        Rental_object.ID_rental_object - 0
        Rental_object.Square           - 1 +
        Rental_object.Description      - 2 +
        Rental_object.Year_built       - 3 + 
        Rental_object.Wall_material    - 4 +
        Rental_object.Floor            - 5 +
        Rental_object.Number_of_floors - 6 +
        Rental_object.Prise            - 7
        Lease_term.Title               - 8
        Object_type.Title              - 9
        Object_state.Title             - 10
        Rental_object.Address          - 11
        Rental_object.ID_client        - 12
        Status_ob.title                - 13
        Rental_object.ID_rental_object - 14
        Rental_object.ID_staff         - 15
        Rental_object.ID_agent         - 16
        Staff.Full_name                - 17
        '''
        '''
        Состав таблицы - Client
        ID_client        - 0
        Passport_data    - 1
        TIN              - 2
        
        Состав таблицы - Natural_person
        Passport_data    - 0
        Surname          - 1
        Name             - 2
        Patronymic       - 3
        Gender           - 4
        Date_of_birth    - 5
        Citizenship      - 6
        
        Состав таблицы - Legal_entity
        TIN              - 0
        Title            - 1
         
        '''

        '''
        Состав процедуры  - show_lease_contract
        
        Search_object.Area             - 0
        Search_object.Square           - 1
        Search_object.Min_price        - 2
        Search_object.Max_price        - 3
        Search_object.Wall_material    - 4
        Search_object.Floor            - 5
        Search_object.Number_of_floor  - 6
        Object_type.Title              - 7
        Search_object.ID_search_object - 8
        Search_object.ID_client        - 9
        Status_ob.title                - 10
        Search_object.ID_staff         - 11
        
        '''

        rental_object_owner = " "
        search_objects_owner = " "

        for contract in lease_contracts:
            for rental_object in rental_objects:
                if contract[7] == rental_object[0]:
                    for client in clients:
                        if rental_object[12] == client[0]:  # Собственики
                            if client[1] == None:  # Дошли до юридических лиц
                                for legal_entity in legal_entitys:
                                    if client[2] == legal_entity[0]:
                                        print("-")
                                        rental_object_owner = legal_entity[1]
                                        pass

                                        '''
                                        all_lease_contract.append(lease_contracts[0],lease_contracts[1],
                                                                  lease_contracts[2],lease_contracts[3],lease_contracts[4],
                                                                  lease_contracts[5],lease_contracts[6],lease_contracts[9],
                                                                  lease_contracts[10],lease_contracts[11],lease_contracts[12],
                                                                  lease_contracts[13],lease_contracts[14],rental_object[9],
                                                                  rental_object[10],legal_entity[1])
                                        '''

                            else:
                                for natural_person in natural_persons:
                                    if client[1] == natural_person[0]:
                                        print("+")
                                        rental_object_owner = natural_person[1] + " " + natural_person[2] + " " + \
                                                              natural_person[3]
                                        pass

                                        '''
                                        all_lease_contract.append(lease_contracts[0], lease_contracts[1],
                                                                  lease_contracts[2], lease_contracts[3],
                                                                  lease_contracts[4], lease_contracts[5], lease_contracts[6],
                                                                  lease_contracts[9], lease_contracts[10], lease_contracts[11],
                                                                  lease_contracts[12], lease_contracts[13], lease_contracts[14],
                                                                  rental_object[9], rental_object[10],
                                                                  )
                                        '''

            for search_object in search_objects:
                if contract[8] == search_object[8]:
                    for client in clients:
                        if search_object[9] == client[0]:
                            if client[1] == None:  # Дошли до юридических лиц

                                for legal_entity in legal_entitys:
                                    if client[2] == legal_entity[0]:
                                        print("=")
                                        search_objects_owner = legal_entity[1]
                            else:
                                for natural_person in natural_persons:
                                    if client[1] == natural_person[0]:
                                        print("--")
                                        search_objects_owner = natural_person[1] + " " + natural_person[2] + " " + \
                                                               natural_person[3]

                            all_lease_contract.append((contract[0], contract[1],
                                                       contract[2], contract[3],
                                                       contract[4], contract[5], contract[6],
                                                       contract[9], contract[10], contract[11],
                                                       contract[12], contract[13], contract[14],
                                                       rental_object[9], rental_object[10], rental_object_owner,
                                                       search_objects_owner
                                                       ))

        cursor.close()
        conn.close()
        print(all_lease_contract)

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return render_template('rent_contract_agent.html', contracts=all_lease_contract)


######################################################################################
#########################          Договоры: Отдел продаж          ###################
######################################################################################

# Cтраница просмотра договоров оказания услуг
@app.route("/rent_contract", methods=['GET', 'POST'])
def show_rent_contract():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_lease_contract')
        lease_contracts = cursor.fetchall()

        cursor.callproc('show_rental_object')
        rental_objects = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        clients = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        legal_entitys = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        natural_persons = cursor.fetchall()

        cursor.callproc('show_search_object')
        search_objects = cursor.fetchall()

        all_lease_contract = []

        '''
        Состав процедуры  - show_lease_contract

        Contract.Number_of_contract    - 0
        Contract.Date                  - 1
        Contract.Amount                - 2
        Rental_object.Address          - 3
        Type_contract.Title            - 4
        Kind_of_contract.Title         - 5
        Contract.ID_staff              - 6
        Contract.ID_rental_object      - 7
        Contract.ID_search_object      - 8
        Rental_object.Square           - 9
        Rental_object.Description      - 10
        Rental_object.Year_built       - 11
        Rental_object.Wall_material    - 12
        Rental_object.Floor            - 13
        Rental_object.Number_of_floors - 14

        '''
        '''

        Состав процедуры  - show_rental_object
        + отмечены поля, которые имеются в show_lease_contract

        Rental_object.ID_rental_object - 0
        Rental_object.Square           - 1 +
        Rental_object.Description      - 2 +
        Rental_object.Year_built       - 3 + 
        Rental_object.Wall_material    - 4 +
        Rental_object.Floor            - 5 +
        Rental_object.Number_of_floors - 6 +
        Rental_object.Prise            - 7
        Lease_term.Title               - 8
        Object_type.Title              - 9
        Object_state.Title             - 10
        Rental_object.Address          - 11
        Rental_object.ID_client        - 12
        Status_ob.title                - 13
        Rental_object.ID_rental_object - 14
        Rental_object.ID_staff         - 15
        Rental_object.ID_agent         - 16
        Staff.Full_name                - 17
        '''
        '''
        Состав таблицы - Client
        ID_client        - 0
        Passport_data    - 1
        TIN              - 2

        Состав таблицы - Natural_person
        Passport_data    - 0
        Surname          - 1
        Name             - 2
        Patronymic       - 3
        Gender           - 4
        Date_of_birth    - 5
        Citizenship      - 6

        Состав таблицы - Legal_entity
        TIN              - 0
        Title            - 1

        '''

        '''
        Состав процедуры  - show_lease_contract

        Search_object.Area             - 0
        Search_object.Square           - 1
        Search_object.Min_price        - 2
        Search_object.Max_price        - 3
        Search_object.Wall_material    - 4
        Search_object.Floor            - 5
        Search_object.Number_of_floor  - 6
        Object_type.Title              - 7
        Search_object.ID_search_object - 8
        Search_object.ID_client        - 9
        Status_ob.title                - 10
        Search_object.ID_staff         - 11

        '''

        rental_object_owner = " "
        search_objects_owner = " "

        for contract in lease_contracts:
            for rental_object in rental_objects:
                if contract[7] == rental_object[0]:
                    for client in clients:
                        if rental_object[12] == client[0]:  # Собственики
                            if client[1] == None:  # Дошли до юридических лиц
                                for legal_entity in legal_entitys:
                                    if client[2] == legal_entity[0]:
                                        print("-")
                                        rental_object_owner = legal_entity[1]
                                        pass

                                        '''
                                        all_lease_contract.append(lease_contracts[0],lease_contracts[1],
                                                                  lease_contracts[2],lease_contracts[3],lease_contracts[4],
                                                                  lease_contracts[5],lease_contracts[6],lease_contracts[9],
                                                                  lease_contracts[10],lease_contracts[11],lease_contracts[12],
                                                                  lease_contracts[13],lease_contracts[14],rental_object[9],
                                                                  rental_object[10],legal_entity[1])
                                        '''

                            else:
                                for natural_person in natural_persons:
                                    if client[1] == natural_person[0]:
                                        print("+")
                                        rental_object_owner = natural_person[1] + " " + natural_person[2] + " " + \
                                                              natural_person[3]
                                        pass

                                        '''
                                        all_lease_contract.append(lease_contracts[0], lease_contracts[1],
                                                                  lease_contracts[2], lease_contracts[3],
                                                                  lease_contracts[4], lease_contracts[5], lease_contracts[6],
                                                                  lease_contracts[9], lease_contracts[10], lease_contracts[11],
                                                                  lease_contracts[12], lease_contracts[13], lease_contracts[14],
                                                                  rental_object[9], rental_object[10],
                                                                  )
                                        '''

            for search_object in search_objects:
                if contract[8] == search_object[8]:
                    for client in clients:
                        if search_object[9] == client[0]:
                            if client[1] == None:  # Дошли до юридических лиц

                                for legal_entity in legal_entitys:
                                    if client[2] == legal_entity[0]:
                                        print("=")
                                        search_objects_owner = legal_entity[1]
                            else:
                                for natural_person in natural_persons:
                                    if client[1] == natural_person[0]:
                                        print("--")
                                        search_objects_owner = natural_person[1] + " " + natural_person[2] + " " + \
                                                               natural_person[3]

                            all_lease_contract.append((contract[0], contract[1],
                                                       contract[2], contract[3],
                                                       contract[4], contract[5], contract[6],
                                                       contract[9], contract[10], contract[11],
                                                       contract[12], contract[13], contract[14],
                                                       rental_object[9], rental_object[10], rental_object_owner,
                                                       search_objects_owner
                                                       ))

        cursor.close()
        conn.close()
        print(all_lease_contract)

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return render_template('rent_contract.html', contracts=all_lease_contract,agent_flag = user_agent)


######################################################################################
#########################          Учет         ######################################
######################################################################################

# Cтраница просмотра объектов поиска
@app.route("/search_objects", methods=['GET', 'POST'])
def show_search_objects():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_search_object')

        r = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        c = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        l = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        n = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_type"))
        rt = cursor.fetchall()

        cursor.close()
        conn.close()

        new_clients = []
        search_ob = []

        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        print(new_clients)

        print(r)
        for row in r:
            for c_row in c:
                if row[9] == c_row[0]:
                    for l_row in l:
                        if c_row[2] == l_row[0]:
                            search_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              l_row[1], row[10]))
                            # row[9] = c_row[1]
                            pass

                    for n_row in n:
                        if c_row[1] == n_row[0]:
                            fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                            search_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              fio_convert, row[10]))
                            # row[9] = fio_convert
                            pass

        print(search_ob)

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    status = []
    status.append(('1', 'Архив'))
    status.append(('2', 'Актуальное'))
    status.append(('4', 'Ожидает оплаты'))

    return render_template('search_objects.html', search_objects=search_ob, clients=new_clients, object_type=rt,
                           status=status,agent_flag = user_agent)

@app.route("/search_wait", methods=['GET', 'POST'])
def show_search_wait():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_search_object')

        r = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        c = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        l = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        n = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_type"))
        rt = cursor.fetchall()

        cursor.close()
        conn.close()

        new_clients = []
        search_ob = []

        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        print(new_clients)

        print(r)
        for row in r:
            for c_row in c:
                if row[9] == c_row[0]:
                    for l_row in l:
                        if c_row[2] == l_row[0]:
                            search_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              l_row[1], row[10]))
                            # row[9] = c_row[1]
                            pass

                    for n_row in n:
                        if c_row[1] == n_row[0]:
                            fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                            search_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              fio_convert, row[10]))
                            # row[9] = fio_convert
                            pass

        print(search_ob)

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    status = []
    status.append(('1', 'Архив'))
    status.append(('2', 'Актуальное'))
    status.append(('4', 'Ожидает оплаты'))

    return render_template('search_wait.html', search_objects=search_ob, clients=new_clients, object_type=rt,
                           status=status,agent_flag = user_agent)

@app.route("/search_objects_archive", methods=['GET', 'POST'])
def show_search_objects_archive():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_search_object')

        r = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        c = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        l = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        n = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_type"))
        rt = cursor.fetchall()

        cursor.close()
        conn.close()

        new_clients = []
        search_ob = []

        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        # print(new_clients)

        # print(r)
        for row in r:
            for c_row in c:
                if row[9] == c_row[0]:
                    for l_row in l:
                        if c_row[2] == l_row[0]:
                            search_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              l_row[1], row[10]))
                            # row[9] = c_row[1]
                            pass

                    for n_row in n:
                        if c_row[1] == n_row[0]:
                            fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                            search_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              fio_convert, row[10]))
                            # row[9] = fio_convert
                            pass

        # print(search_ob)

        status = []
        status.append(('1', 'Архив'))
        status.append(('2', 'Актуальное'))
        status.append(('4', 'Ожидает оплаты'))


    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return render_template('search_objects_archive.html', search_objects=search_ob, clients=new_clients, object_type=rt,
                           status=status,agent_flag = user_agent)

@app.route('/search_objects/add', methods=["POST"])
def search_objects_add():
    try:
        print("Зашли")
        area = request.form['area']
        square = request.form['square']
        min_price = request.form['min_price']
        max_price = request.form['max_price']
        wall_material = request.form['wall_material']
        floor = request.form['floor']
        number_of_floor = request.form['number_of_floor']
        id_client = request.form['id_client']
        id_object_type = request.form['id_object_type']
        status = request.form['status_type']

        conn = mysql.connect()
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO Search_object (Area, Square, Min_price, Max_price, Wall_material, Floor, Number_of_floor, ID_client, ID_object_type, ID_status) VALUES (%s, %s,%s,%s,%s,%s,%s,%s,%s,%s)",
            (area, square, min_price, max_price, wall_material, floor, number_of_floor, id_client, id_object_type, status))
        conn.commit()
        print("Вышли")
    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return redirect(url_for('show_sales'))

@app.route('/search_objects/update', methods=["POST"])
def search_objects_update():
    area = request.form['area']
    square = request.form['square']
    min_price = request.form['min_price']
    max_price = request.form['max_price']
    wall_material = request.form['wall_material']
    floor = request.form['floor']
    number_of_floor = request.form['number_of_floor']
    id_client = request.form['id_client']
    id_object_type = request.form['id_object_type']
    id_search = request.form['id']
    status = request.form['status_type']

    conn = mysql.connect()
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE Search_object SET Area =%s, Square=%s, Min_price=%s, Max_price=%s, Wall_material=%s,Floor=%s ,Number_of_floor=%s ,ID_client=%s ,ID_object_type=%s, ID_status=%s WHERE ID_search_object=%s",
        (area, square, min_price, max_price, wall_material, floor, number_of_floor, id_client, id_object_type, status,
         id_search))
    conn.commit()

    return redirect(url_for('show_search_objects'))

@app.route('/search_objects/delete/<string:id_data>', methods=["GET"])
def search_objects_delete(id_data):
    conn = mysql.connect()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM Search_object WHERE ID_search_object=%s", (id_data,))
    conn.commit()
    return redirect(url_for('show_search_objects'))


# Cтраница просмотра объектов аренды
@app.route("/rental_objects", methods=['GET', 'POST'])
def show_rental_objects():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_rental_object')

        r = cursor.fetchall()

        cursor.execute(("SELECT * FROM Lease_term"))
        le = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_type"))
        o = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_state"))
        os = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        c = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        l = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        n = cursor.fetchall()

        new_clients = []
        rental_ob = []

        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        #print(new_clients)

        #print(r)
        for row in r:
            for c_row in c:
                if row[12] == c_row[0]:
                    for l_row in l:
                        if c_row[2] == l_row[0]:
                            rental_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              row[9], row[10], row[11], l_row[1], row[13]))
                            # row[9] = c_row[1]
                            pass

                    for n_row in n:
                        if c_row[1] == n_row[0]:
                            fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                            rental_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              row[9], row[10], row[11], fio_convert, row[13]))
                            # row[9] = fio_convert
                            pass



        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    status = []
    status.append(('1', 'Архив'))
    status.append(('2', 'Актуальное'))
    status.append(('3', 'Сдается'))

    print(rental_ob)
    print(user_agent)

    return render_template('rental_objects.html', rental_objects=rental_ob, lease_term=le, object_type=o,
                           object_state=os, status=status, clients=new_clients,agent_flag = user_agent)

@app.route("/rental_archive", methods=['GET', 'POST'])
def show_rental_archive():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_rental_object')

        r = cursor.fetchall()

        cursor.execute(("SELECT * FROM Lease_term"))
        le = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_type"))
        o = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_state"))
        os = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        c = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        l = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        n = cursor.fetchall()

        new_clients = []
        rental_ob = []

        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        #print(new_clients)

        #print(r)
        for row in r:
            for c_row in c:
                if row[12] == c_row[0]:
                    for l_row in l:
                        if c_row[2] == l_row[0]:
                            rental_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              row[9], row[10], row[11], l_row[1], row[13]))
                            # row[9] = c_row[1]
                            pass

                    for n_row in n:
                        if c_row[1] == n_row[0]:
                            fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                            rental_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              row[9], row[10], row[11], fio_convert, row[13]))
                            # row[9] = fio_convert
                            pass



        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    status = []
    status.append(('1', 'Архив'))
    status.append(('2', 'Актуальное'))
    status.append(('3', 'Сдается'))
    status.append(('4', 'Ожидает оплаты'))

    print(rental_ob)

    return render_template('rental_archive.html', rental_objects=rental_ob, lease_term=le, object_type=o,
                           object_state=os, status=status, clients=new_clients, agent_flag = user_agent)

@app.route("/rental_wait", methods=['GET', 'POST'])
def show_rental_wait():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_rental_object')

        r = cursor.fetchall()

        cursor.execute(("SELECT * FROM Lease_term"))
        le = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_type"))
        o = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_state"))
        os = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        c = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        l = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        n = cursor.fetchall()

        new_clients = []
        rental_ob = []

        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        #print(new_clients)

        #print(r)
        for row in r:
            for c_row in c:
                if row[12] == c_row[0]:
                    for l_row in l:
                        if c_row[2] == l_row[0]:
                            rental_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              row[9], row[10], row[11], l_row[1], row[13]))
                            # row[9] = c_row[1]
                            pass

                    for n_row in n:
                        if c_row[1] == n_row[0]:
                            fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                            rental_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              row[9], row[10], row[11], fio_convert, row[13]))
                            # row[9] = fio_convert
                            pass



        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    status = []
    status.append(('1', 'Архив'))
    status.append(('2', 'Актуальное'))
    status.append(('3', 'Сдается'))
    status.append(('4', 'Ожидает оплаты'))

    print(rental_ob)

    return render_template('rental_wait.html', rental_objects=rental_ob, lease_term=le, object_type=o,
                           object_state=os, status=status, clients=new_clients, agent_flag = user_agent)

@app.route("/rental_rent", methods=['GET', 'POST'])
def show_rental_rent():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_rental_object')

        r = cursor.fetchall()

        cursor.execute(("SELECT * FROM Lease_term"))
        le = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_type"))
        o = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_state"))
        os = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        c = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        l = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        n = cursor.fetchall()

        new_clients = []
        rental_ob = []

        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        #print(new_clients)

        #print(r)
        for row in r:
            for c_row in c:
                if row[12] == c_row[0]:
                    for l_row in l:
                        if c_row[2] == l_row[0]:
                            rental_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              row[9], row[10], row[11], l_row[1], row[13]))
                            # row[9] = c_row[1]
                            pass

                    for n_row in n:
                        if c_row[1] == n_row[0]:
                            fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                            rental_ob.append((row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                              row[9], row[10], row[11], fio_convert, row[13]))
                            # row[9] = fio_convert
                            pass



        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    status = []
    status.append(('1', 'Архив'))
    status.append(('2', 'Актуальное'))
    status.append(('3', 'Сдается'))
    status.append(('4', 'Ожидает оплаты'))

    print(rental_ob)

    return render_template('rental_rent.html', rental_objects=rental_ob, lease_term=le, object_type=o,
                           object_state=os, status=status, clients=new_clients,agent_flag = user_agent)

@app.route('/rental_objects/add', methods=["POST"])
def rental_objects_add():
    try:
        print("Зашли")
        square = request.form['square']
        description = request.form['description']
        year_built = request.form['year_built']
        wall_material = request.form['wall_material']
        floor = request.form['floor']
        number_of_floors = request.form['number_of_floors']
        prise = request.form['prise']
        id_lease_term = request.form['id_lease_term']
        id_object_type = request.form['id_object_type']
        id_object_state = request.form['id_object_state']
        id_client = request.form['id_client']
        address = request.form['address']
        id_agent = request.form['id_agent']
        print("Агент")
        print(id_agent)
        id_status = '4'

        global user_sales
        conn = mysql.connect()
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO Rental_object (Square, Description, Year_built, Wall_material, Floor, Number_of_floors, Prise, ID_lease_term, ID_object_type, ID_object_state, ID_client, Address, ID_status,ID_staff,ID_agent) VALUES (%s, %s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
            (square, description, year_built, wall_material, floor, number_of_floors, prise, id_lease_term, id_object_type,
             id_object_state, id_client, address, id_status,user_sales,id_agent))
        conn.commit()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    return redirect(url_for('show_sales'))

@app.route('/rental_objects/update', methods=["POST"])
def rental_objects_update():
    print("Зашли")
    id_rental_object = request.form['id_rental_object']
    square = request.form['square']
    description = request.form['description']
    year_built = request.form['year_built']
    wall_material = request.form['wall_material']
    floor = request.form['floor']
    number_of_floors = request.form['number_of_floors']
    prise = request.form['prise']
    id_lease_term = request.form['id_lease_term']
    print(id_lease_term)
    id_object_type = request.form['id_object_type']
    id_object_state = request.form['id_object_state']
    id_client = request.form['id_client']
    address = request.form['address']
    id_status = request.form['id_status']
    global user_sales

    conn = mysql.connect()
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE Rental_object SET Square=%s, Description=%s, Year_built=%s, Wall_material=%s, Floor=%s, Number_of_floors=%s, Prise=%s, ID_lease_term=%s, ID_object_type=%s, ID_object_state=%s, ID_client=%s, Address=%s, ID_status=%s, ID_staff=%s WHERE ID_rental_object=%s",
        (square, description, year_built, wall_material, floor, number_of_floors, prise, id_lease_term, id_object_type,
         id_object_state, id_client, address, id_status,user_sales, id_rental_object))
    conn.commit()

    return redirect(url_for('show_rental_objects'))

@app.route('/rental_objects/delete/<string:id_data>', methods=["GET"])
def rental_objects_delete(id_data):
    conn = mysql.connect()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM Rental_object WHERE ID_rental_object=%s", (id_data,))
    conn.commit()
    return redirect(url_for('show_rental_objects'))

# Cтраница просмотра продаж
@app.route("/sales", methods=['GET', 'POST'])
def show_sales():
    global user_sales
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM Staff_sales_plan_temp")
        re = cursor.fetchall()

        plan1 = []

        for r in re:
            t = str(r[3])
            if (t == user_sales):
                plan1 = r

        cursor.callproc('show_rental_object')
        r = cursor.fetchall()

        cursor.execute(("SELECT * FROM Lease_term"))
        lease = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_type"))
        o = cursor.fetchall()

        cursor.execute(("SELECT * FROM Object_state"))
        os = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        c = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        l = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        n = cursor.fetchall()

        new_clients = []
        rental_ob = []

        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        for row in r:
            for c_row in c:
                if row[12] == c_row[0]:
                    for l_row in l:
                        if c_row[2] == l_row[0]:
                            t = str(row[15])
                            if (t == user_sales):
                                rental_ob.append(
                                    (row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                     row[9], row[10], row[11], l_row[1], row[13]))
                            pass

                    for n_row in n:
                        if c_row[1] == n_row[0]:

                            t = str(row[15])
                            if (t == user_sales):

                                fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                                rental_ob.append(
                                    (row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                     row[9], row[10], row[11], fio_convert, row[13]))
                            pass


        cursor.callproc('show_search_object')
        r = cursor.fetchall()
        print("процедура")
        print(r)

        search_ob = []
        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        for row in r:
            for c_row in c:
                if row[9] == c_row[0]:
                    for l_row in l:
                        if c_row[2] == l_row[0]:
                            t = str(row[11])
                            if (t == user_sales):
                                search_ob.append(
                                    (row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                     l_row[1], row[10]))

                            pass

                    for n_row in n:
                        if c_row[1] == n_row[0]:
                            t = str(row[11])

                            if (t == user_sales):
                                fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                                search_ob.append(
                                    (row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8],
                                     fio_convert, row[10]))

                            pass

        cursor.callproc('login')
        login = cursor.fetchall()

        cursor.close()
        conn.close()

    except Exception as e:
        print("Поподаем в exception")
        print(e)

    print("Серч")
    print(search_ob)

    status = []
    status.append(('1', 'Архив'))
    status.append(('2', 'Актуальное'))

    sort_renatl_objects = sort_rent(rental_ob)

    return render_template('sales.html', plan=plan1, rental_objects=sort_renatl_objects, search_objects=search_ob,
                           clients=new_clients, lease_term=lease, object_type=o, object_state=os, agent=login,status = status)


######################################################################################
#########################          Аналитика         #################################
######################################################################################



#Cтраница просмотра платежей за услуги
@app.route("/payments_service", methods=['GET', 'POST'])
def show_payments_service():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_payments_service')
        r = cursor.fetchall()

        cursor.close()
        conn.close()
    except Exception as e:
        print("Попадаем в exception")
        print(e)
    return render_template('payments_service.html', payments_service=r)

#Cтраница просмотра платежей за услуги тестовая версия
@app.route("/payments_service_new", methods=['GET', 'POST'])
def show_payments_service_new():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_payments_service')
        r = cursor.fetchall()


        cursor.close()
        conn.close()

    except Exception as e:
        print("Попадаем в exception")
        print(e)

    status = []
    status.append(('2', 'Актуальное'))
    status.append(('4', 'Ожидает оплаты'))



    return render_template('payments_service_new.html', payments_service_new=r, status=status, agent_flag = user_agent)


@app.route('/payments_service_new/update', methods=["POST"])
def payments_service_new_update():

    number1 = request.form['number1']
    id_status_ob = request.form['id_status_ob']
    print("Вошли")
    print(id_status_ob)

    conn = mysql.connect()
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE Contract SET  ID_status_ob=%s WHERE Number_of_contract=%s",
        (id_status_ob,number1))
    conn.commit()

    return redirect(url_for('show_payments_service_new'))

#Cтраница просмотра платежей по договору
@app.route("/payments_receipt_contract", methods=['GET', 'POST'])
def show_payments_receipt_contract():
    try:
        conn = mysql.connect()
        cursor = conn.cursor()

        cursor.callproc('show_payments_receipt_contract')
        r = cursor.fetchall()

        cursor.execute(("SELECT * FROM Client"))
        c = cursor.fetchall()

        cursor.execute(("SELECT * FROM Legal_entity"))
        l = cursor.fetchall()

        cursor.execute(("SELECT * FROM Natural_person"))
        n = cursor.fetchall()

        new_clients = []

        fio_convert = ''

        for row in c:
            for l_row in l:
                if row[2] == l_row[0]:
                    new_clients.append((row[0], l_row[1]))
                    pass

            for n_row in n:
                if row[1] == n_row[0]:
                    fio_convert = n_row[1] + " " + n_row[2] + " " + n_row[3]
                    new_clients.append((row[0], fio_convert))
                    pass

        print(new_clients)



        cursor.close()
        conn.close()

    except Exception as e:
        print("Попадаем в exception")
        print(e)



    return render_template('payments_receipt_contract.html', payments_receipt_contract=r, agent_flag = user_agent, clients=new_clients)


######################################################################################
#########################          Поиск         #####################################
######################################################################################

###############################################################################################
#########################          Формирование договора         ##############################
###############################################################################################
@app.route('/file_downloads_services', methods=["POST"])
def file_downloads():
    number_dogovor = request.form['number_of_dogovor']
    date1 = request.form['date_of_rent']
    year = date1[:4]
    month = date1[5:7]
    day = date1[8:10]
    owner = request.form['owner']
    name = request.form['name']
    document = Document('services_contract_template.docx')
    for p in document.paragraphs:
        p.text = re.sub(r'НОМЕР КОНТРАКТА', number_dogovor, p.text)
        p.text = re.sub(r'ДЕНЬ', day, p.text)
        p.text = re.sub(r'МЕСЯЦ', month, p.text)
        p.text = re.sub(r'ГОД', year, p.text)
        p.text = re.sub(r'ЗАКАЗЧИК', owner, p.text)
        p.text = re.sub(r'НАИМЕНОВАНИЕ', name, p.text)
    document.save('services_contract_downloads.docx')
    return send_file('services_contract_downloads.docx', attachment_filename='services_contract_downloads.docx')


@app.route('/file_downloads_rent', methods=["POST"])
def file_downloads_rent():
    number_dogovor = request.form['number_of_dogovor']
    date1 = request.form['date_of_rent']
    year = date1[:4]
    month = date1[5:7]
    day = date1[8:10]
    landlord = request.form['landlord']
    tenant = request.form['tenant']
    address = request.form['address']
    srok = request.form['srok']
    price = request.form['price']

    document = Document('rent_contract_template.docx')
    for p in document.paragraphs:
        p.text = re.sub(r'НОМЕР АРЕНДЫ', number_dogovor, p.text)
        p.text = re.sub(r'ДЕНЬ АРЕНДЫ', day, p.text)
        p.text = re.sub(r'МЕСЯЦ АРЕНДЫ', month, p.text)
        p.text = re.sub(r'ГОД АРЕНДЫ', year, p.text)
        p.text = re.sub(r'АРЕНДОДАТЕЛЬ', landlord, p.text)
        p.text = re.sub(r'АРЕНДАТОР', tenant, p.text)
        p.text = re.sub(r'АДРЕС', address, p.text)
        p.text = re.sub(r'СРОК ', srok, p.text)
        p.text = re.sub(r'СУММА ', price, p.text)

    document.save('rent_contract_downloads.docx')
    return send_file('rent_contract_downloads.docx', attachment_filename='rent_contract_downloads.docx')

@app.route('/file_downloads_lease', methods=["POST"])
def file_downloads_lease():
    return send_file('services_contract.docx', attachment_filename='services_contract.docx')

@app.route('/file_downloads_rent_temp', methods=["POST"])
def file_downloads_rent_temp():
    return send_file('rent_contract.docx', attachment_filename='rent_contract.docx')

if __name__ == '__main__':
    app.run()
